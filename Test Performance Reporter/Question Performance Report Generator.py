from docx import Document
import csv

path = 'C:\\Users\\ALLAN\\Desktop\\Computers\\Coding\\Python\\__1.   Projects\\Personal Projects\\_Public\\Projects\\PSAT - Word Compiler\\'
section3 = csv.reader(open(path + 'Student Exam Report - @3) 2018 PSAT Math (No Calculator) - Student Exam Report.csv'))
section4 = csv.reader(open(path + 'Student Exam Report - @4) 2018 PSAT Math (Calculator) - Student Exam Report.csv'))

section3_AnswerKey = [['B'], ['D'], ['D'], ['A'], ['D'], ['D'], ['D'], ['A'], ['B'], ['C'], ['B'], ['C'], ['C'], ['1'], ['81'], ['9'], ['4/3']]
section4_AnswerKey = [['B'], ['D'], ['A'], ['D'], ['D'], ['C'], ['D'], ['B'], ['C'], ['C'], ['B'], ['D'], ['C'], ['A'], ['C'], ['D'], ['C'], ['D'], ['C'], ['B'], ['B'], ['D'], ['D'], ['D'], ['C'], ['D'], ['A'], ['540'], ['1.21'], ['1.13'], ['166']]

student = dict()
#student[id] = [Section3, Section4]
#Section3 = [Name, Points, Score, [Q1, Q2, Q3, ....]]
'''
for row in section3:
    if row[0] == 'GradeCam ID':
        continue
    id = row[0]
    Name = ''
    Points = ''
    Score = ''
    answers = list()
    for index in range(1, len(row)):
        if index == 1:
            Name = row[index]
        elif index == 2:
            Points = row[index]
        elif index == 3:
            Score = row[index]
        else:
            answers.append(row[index])

    Section3List = [Name, Points, Score, answers]
    student[id] = [Section3List, []]

for row in section4:
    if row[0] == 'GradeCam ID':
        continue
    id = row[0]
    #if id not in student:
    #    print('ERROR MISSING ID:' + str(id))
    Name = ''
    Points = ''
    Score = ''
    answers = list()
    for index in range(1, len(row)):
        if index == 1:
            Name = row[index]
        elif index == 2:
            Points = row[index]
        elif index == 3:
            Score = row[index]
        else:
            answers.append(row[index])

    Section4List = [Name, Points, Score, answers]
    if id in student:
        student[id] = [student[id][0], Section4List]
    else:
        student[id] = [[], Section4List]
'''
document = Document()

Sec3 = dict()
i = -1
for row in section3:
    for index in range(4, len(row)):
        if row[0] == 'GradeCam ID':
            Sec3['Question ' + str(index-3)] = 0
            continue
        if row[index] in section3_AnswerKey[index-4]:
            Sec3['Question ' + str(index-3)] +=1
    i+=1

for question in Sec3:
    Sec3[question] = -Sec3[question]/float(i) *100

ordered_difficulty_sec3 = sorted(Sec3.items(), reverse=True, key=lambda kv: kv[1])


Sec4 = dict()
i = -1
for row in section4:
    for index in range(4, len(row)):
        if row[0] == 'GradeCam ID':
            Sec4['Question ' + str(index-3)] = 0
            continue
        if row[index] in section4_AnswerKey[index-4]:
            Sec4['Question ' + str(index-3)] +=1
    i+=1

for question in Sec4:
    Sec4[question] = -Sec4[question]/float(i) *100

ordered_difficulty_sec4 = sorted(Sec4.items(), reverse=True, key=lambda kv: kv[1])

document.add_heading('Section 3', level=1)
for question in ordered_difficulty_sec3:
    document.add_paragraph(question[0] + ': ' + str(int(-Sec3[question[0]])) + '%', style='List Bullet')

document.add_page_break()

document.add_heading('Section 4', level=1)
for question in ordered_difficulty_sec4:
    document.add_paragraph(question[0] + ': ' + str(int(-Sec4[question[0]])) + '%', style='List Bullet')

document.save(path + 'Section 3-4 Question Performance Report.docx')
